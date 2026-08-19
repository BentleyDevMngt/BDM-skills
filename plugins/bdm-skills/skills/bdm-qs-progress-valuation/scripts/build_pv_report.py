#!/usr/bin/env python3
"""
Build a BDM QS Monthly Progress Valuation Report (Form 425) from the latest template
working copy and a JSON content map. Deterministic build only — Claude authors the map.

Faithful population: every populated cell PRESERVES the template cell's native formatting
(font / size / alignment), so the output matches the baseline template. No flags, no comments.

Phase 1 (python-docx): token fill, table cells, narrative, trade breakdown, variations,
  cashflow table, consultants register, drafting-note strip, row trims, expired-date bold-red,
  photo captions, appendix-marker cleanup, DRAFT marker.
Phase 2 (zip/XML): site-photo image swap, cashflow chart cache, header/footer tokens.
(--pdf): render the docx to PDF; if `compile_appendices` is set, append the contractor's
  claim + a Statutory-Declaration note after their Appendix dividers.

Usage:
  python build_pv_report.py --template 425.docx --config map.json \
      [--photos onenote.pdf|folder] [--claimdir "<period>/Builder Claim"] --out OUT.docx [--pdf]
"""
import argparse, json, copy, os, re, zipfile, glob, datetime, subprocess, tempfile
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setpar(p, text):
    rpr=None
    if p.runs:
        e=p.runs[0]._r.find(qn('w:rPr'))
        if e is not None: rpr=copy.deepcopy(e)
    for r in list(p._p.findall(qn('w:r'))): p._p.remove(r)
    nr=OxmlElement('w:r')
    if rpr is not None: nr.append(rpr)
    t=OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text=text; nr.append(t); p._p.append(nr)

def set_cell(d, ti, r, c, text, bold=False, red=False):
    """Populate a cell, PRESERVING the template cell's native run formatting (font/size/etc.)
    so the output does not deviate from the baseline template."""
    cell=d.tables[ti].cell(r,c); p=cell.paragraphs[0]
    rpr=None
    if p.runs:
        e=p.runs[0]._r.find(qn('w:rPr'))
        if e is not None: rpr=copy.deepcopy(e)
    for x in cell.paragraphs[1:]: x._p.getparent().remove(x._p)
    for r_ in list(p._p.findall(qn('w:r'))): p._p.remove(r_)
    nr=OxmlElement('w:r')
    if rpr is not None: nr.append(rpr)
    tnode=OxmlElement('w:t'); tnode.set(qn('xml:space'),'preserve'); tnode.text=str(text)
    nr.append(tnode); p._p.append(nr)
    if bold or red:
        run=docx.text.run.Run(nr, p)
        if bold: run.bold=True
        if red: run.font.color.rgb=RGBColor(0xC0,0x00,0x00); run.bold=True
    return cell

def del_row(d, ti, r):
    tbl=d.tables[ti]; tbl._tbl.remove(tbl.rows[r]._tr)

def parse_date(s):
    for fmt in ('%d/%m/%Y','%d %B %Y','%d %b %Y','%Y-%m-%d'):
        try: return datetime.datetime.strptime(s.strip(),fmt).date()
        except: pass
    return None

def add_additional_appendices(d, cfg, claimdir):
    """Append optional standalone appendices (heading + embedded PDF page-images) before the
    document's final section break. Use for content not always present (e.g. development-cost
    invoices)."""
    apps=cfg.get('additional_appendices')
    if not apps: return
    from docx.text.paragraph import Paragraph
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    body=d.element.body; sect=body.find(qn('w:sectPr'))
    parent=d.paragraphs[0]._parent
    def new_para():
        p=OxmlElement('w:p')
        (sect.addprevious(p) if sect is not None else body.append(p))
        return Paragraph(p, parent)
    for app in apps:
        h=new_para()
        try: h.style=d.styles['Heading 2']
        except Exception: pass
        h.add_run(app['heading']); h.paragraph_format.page_break_before=True
        if app.get('intro'):
            new_para().add_run(app['intro'])
        for f in app.get('pdfs',[]):
            path=f if os.path.isabs(f) else os.path.join(claimdir or '', f)
            if not os.path.exists(path): continue
            for fp,w,hh in rasterize_pdf(path):
                pp=new_para(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                run=pp.add_run(); iw=6.3; ih=hh*(iw/w)
                if ih>9.2: iw=iw*(9.2/ih)
                run.add_picture(fp, width=Inches(iw))

def clean_appendix_markers(d):
    """Remove any @@APPENDIX_B markers left in the template body."""
    for p in list(d.paragraphs):
        if '@@APPENDIX_B' in p.text:
            p._p.getparent().remove(p._p)

def fallback_token_replace(d, TOK):
    """Catch tokens split across runs (run-level replace misses them): join each paragraph's
    runs, replace, and rewrite only if a token was present. Fixes cover Valuation No./Report Date."""
    def fix(paras):
        for p in paras:
            full=''.join(r.text for r in p.runs)
            if any(k in full for k in TOK):
                newt=full
                for k,v in TOK.items(): newt=newt.replace(k,v)
                if newt!=full: setpar(p, newt)
    fix(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                fix(cell.paragraphs)

def rasterize_pdf(path, dpi=140):
    import fitz
    doc=fitz.open(path); tmp=tempfile.mkdtemp(); out=[]
    for i,pg in enumerate(doc):
        pix=pg.get_pixmap(dpi=dpi)
        fp=os.path.join(tmp,'p_%03d.png'%i); pix.save(fp); out.append((fp,pix.width,pix.height))
    return out

def _img_para(after_el, parent, fp, w, h, max_w=6.3):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    newp=OxmlElement('w:p'); after_el.addnext(newp)
    para=docx.text.paragraph.Paragraph(newp, parent)
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after=Pt(4); para.paragraph_format.space_before=Pt(4)
    run=para.add_run(); iw=max_w; ih=h*(iw/w)
    if ih>9.2: iw=iw*(9.2/ih)
    run.add_picture(fp, width=Inches(iw))
    return newp

def embed_appendices(d, cfg, claimdir):
    """Embed annexure PDFs as page images in the body, like the template."""
    ap=cfg.get('embed_appendices')
    if not ap: return
    start=end=None
    for p in d.paragraphs:
        if '@@APPENDIX_B_START@@' in p.text: start=p
        if '@@APPENDIX_B_END@@' in p.text: end=p
    if start is not None and end is not None:
        el=start._p.getnext()
        while el is not None and el is not end._p:
            nxt=el.getnext(); el.getparent().remove(el); el=nxt
        anchor=start._p
        for f in ap.get('appendix_b',[]):
            path=f if os.path.isabs(f) else os.path.join(claimdir or '', f)
            if not os.path.exists(path): continue
            for fp,w,h in rasterize_pdf(path):
                anchor=_img_para(anchor, start._parent, fp, w, h)
        start._p.getparent().remove(start._p); end._p.getparent().remove(end._p)
    cnote=ap.get('appendix_c')
    if cnote:
        ph=None
        for p in d.paragraphs:
            if 'Replace the page below' in p.text: ph=p; break
        if ph is not None:
            if isinstance(cnote,dict) and cnote.get('pdf'):
                path=cnote['pdf'] if os.path.isabs(cnote['pdf']) else os.path.join(claimdir or '',cnote['pdf'])
                anchor=ph._p
                if os.path.exists(path):
                    for fp,w,h in rasterize_pdf(path):
                        anchor=_img_para(anchor, ph._parent, fp, w, h)
                ph._p.getparent().remove(ph._p)
            else:
                setpar(ph, cnote.get('note') if isinstance(cnote,dict) else str(cnote))

def set_cell_obj(cell, text, bold=None):
    p=cell.paragraphs[0]; rpr=None
    if p.runs:
        e=p.runs[0]._r.find(qn('w:rPr'))
        if e is not None: rpr=copy.deepcopy(e)
    for x in cell.paragraphs[1:]: x._p.getparent().remove(x._p)
    for r_ in list(p._p.findall(qn('w:r'))): p._p.remove(r_)
    nr=OxmlElement('w:r')
    if rpr is not None: nr.append(rpr)
    tnode=OxmlElement('w:t'); tnode.set(qn('xml:space'),'preserve'); tnode.text=str(text)
    nr.append(tnode); p._p.append(nr)
    if bold is not None:
        from docx.text.run import Run
        Run(nr,p).bold=bold

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),color); tcPr.append(sh)

def bold_shade_row(d, ti, ri, fill='F4F4F1'):
    for cell in d.tables[ti].rows[ri].cells:
        for p in cell.paragraphs:
            for r in p.runs: r.bold=True
        shade(cell, fill)

def justify_body(d):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for p in d.paragraphs:
        if p.style and p.style.name=='Normal' and p.text.strip():
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def apply_row_format(d, cfg):
    # font size (pt) for every cell of named tables
    for spec in cfg.get('cell_size',[]):
        ti=spec['table']; pt=spec['size_pt']
        for row in d.tables[ti].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size=Pt(pt)
    # explicit per-row bold: BOLD the listed rows, regular for all others (handles merged cells)
    for spec in cfg.get('row_bold',[]):
        ti=spec['table']; bold_rows=set(spec['bold_rows']); t=d.tables[ti]
        for ri in range(len(t.rows)):
            want=ri in bold_rows
            for cell in t.rows[ri].cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.bold=want

def clone_2col_table(d, cfg):
    spec=cfg.get('dev_cost_table')
    if not spec: return
    import copy as _c
    from docx.table import Table
    anchor=None
    for p in d.paragraphs:
        if spec['anchor'] in p.text: anchor=p; break
    if anchor is None: return
    model=d.tables[spec['model_table']]._tbl
    new=_c.deepcopy(model); anchor._p.addnext(new)
    t=Table(new, anchor._p.getparent())
    rows=spec['rows']
    clean=_c.deepcopy(t.rows[1]._tr)            # a clean data row from the model
    for _ in range(len(t.rows)-2):              # remove ALL existing data rows (avoid inheriting styled rows)
        t._tbl.remove(t.rows[1]._tr)
    for _ in range(len(rows)):                  # insert clean data rows before the (styled) total row
        t.rows[len(t.rows)-1]._tr.addprevious(_c.deepcopy(clean))
    allrows=[spec['header']]+rows+[spec['total']]
    for ri,rv in enumerate(allrows):
        for ci,val in enumerate(rv):
            set_cell_obj(t.cell(ri,ci), val)
    return t

def cleanup_blank_paras(d):
    """Collapse consecutive runs-less empty paragraphs (kills stray blank pages), but NEVER
    remove a paragraph carrying a section break or page break."""
    prev_empty=False
    for p in list(d.paragraphs):
        xml=p._p.xml
        keep = ('<w:drawing' in xml) or ('<w:sectPr' in xml) or ('<w:br' in xml) or ('pageBreakBefore' in xml)
        is_empty = (not p.text.strip()) and (not keep)
        if is_empty and prev_empty: p._p.getparent().remove(p._p)
        prev_empty=is_empty

def _money(v):
    s=str(v).strip()
    if s=='' or s in ('-','—','–'): return 0.0
    neg = s.lstrip().startswith('(') or s.lstrip().startswith('-')
    s=re.sub(r'[^0-9.]','',s)
    if s in ('','.'): return None
    try: x=float(s)
    except: return None
    return -x if neg else x

def apply_inserts(d, cfg):
    for spec in cfg.get('insert_rows',[]):
        ti=spec['table']; after=spec['after_row']; n=spec.get('count',1)
        t=d.tables[ti]
        for _ in range(n):
            t.rows[after]._tr.addnext(copy.deepcopy(t.rows[after]._tr))

def apply_bold_shade(d, cfg):
    for spec in cfg.get('bold_shade_rows',[]):
        ti=spec['table']
        for ri in spec['rows']:
            if ri < len(d.tables[ti].rows): bold_shade_row(d, ti, ri)

def qa_checks(docx_path, cfg):
    """Internal-consistency + carry-forward QA. Returns list of (name, ok, got, expect, delta)."""
    qd=docx.Document(docx_path); checks=[]
    def chk(name, got, expect, tol=1.0):
        ok = (got is not None and expect is not None and abs(got-expect)<=tol)
        checks.append((name, ok, got, expect, (None if (got is None or expect is None) else round(got-expect,2))))
    def find(pred):
        for t in qd.tables:
            try:
                if pred(t): return t
            except: pass
        return None
    cert=find(lambda t: 'PROGRESS CERTIFICATE' in t.cell(0,0).text)
    C={}
    if cert:
        for r in range(len(cert.rows)):
            C[cert.rows[r].cells[0].text.strip()]=_money(cert.rows[r].cells[-1].text)
    g=lambda pfx: next((v for k,v in C.items() if k.startswith(pfx)), None)
    if cert:
        chk('Cert: Adjusted = Original + Approved variations', g('Adjusted Contract Sum'),(g('Original Contract Sum') or 0)+(g('Approved Variations') or 0))
        chk('Cert: Total works = Contract works + Variation works', g('Total Value of Works Completed'),(g('Contract Works Complete') or 0)+(g('Variation Works Complete') or 0))
        chk('Cert: Net this claim = Net value to date − Previous recommendation', g('Net Sum Recommended This Claim'),(g('Net Value of Payment to Date') or 0)-(g('Less: Previous Net Recommendation') or 0))
        chk('Cert: GST = 10% of net this claim', g('GST'), round((g('Net Sum Recommended This Claim') or 0)*0.10,2))
        chk('Cert: Net incl GST = net × 1.10', g('NET SUM INCL GST'), round((g('Net Sum Recommended This Claim') or 0)*1.10,2))
    # trade subtotal = certified contract works
    tb=find(lambda t: len(t.rows)>20 and t.cell(len(t.rows)-1,0).text.strip().lower()=='subtotal')
    if tb and cert:
        sub=_money(tb.cell(len(tb.rows)-1,len(tb.columns)-2).text)  # Total $ column (second last)
        chk('Trade-breakdown subtotal = Contract works complete (cert)', sub, g('Contract Works Complete'))
    # carry-forward vs prior report
    prior=cfg.get('prior',{})
    if cert and 'net_value_to_date' in prior:
        chk('Carry-fwd: Previous recommendation = prior Net value to date', g('Less: Previous Net Recommendation'), prior['net_value_to_date'])
    # §2.4 construction costs previously-assessed = prior total completed
    cc=find(lambda t: t.cell(0,0).text.strip()=='Description' and 'Construction Works' in t.cell(1,0).text)
    if cc and 'construction_prev_assessed' in prior:
        chk('Carry-fwd: Construction prev-assessed = prior total completed', _money(cc.cell(1,2).text), prior['construction_prev_assessed'])
    # professional fees previously-claimed line
    if 'prof_fees_to_date' in prior:
        for t in qd.tables:
            for r in range(len(t.rows)):
                if 'previously claimed to date' in t.cell(r,0).text.lower():
                    chk('Carry-fwd: Prof fees previously claimed = prior aggregate', _money(t.cell(r,len(t.columns)-1).text), prior['prof_fees_to_date']); break
    # §2.6 Contingency self-reconciliation: drawdown = Σ(Less rows); remaining = contingency − drawdown
    cont=find(lambda t: len(t.columns)==2 and any('contingency' in t.cell(r,0).text.lower() and 'building' in t.cell(r,0).text.lower() for r in range(len(t.rows))))
    if cont:
        carried=less=draw=remain=None
        for r in range(len(cont.rows)):
            lbl=cont.cell(r,0).text.lower(); amt=_money(cont.cell(r,1).text)
            if 'contingency' in lbl and 'building' in lbl and carried is None: carried=amt
            elif lbl.startswith('less'): less=(less or 0)+abs(amt or 0)
            elif 'total' in lbl and 'drawdown' in lbl: draw=abs(amt or 0)
            elif 'remaining' in lbl: remain=amt
        if draw is not None and less is not None:
            chk('Contingency: Total drawdown = Σ(Less variation/overrun rows)', draw, less)
        if remain is not None and carried is not None and draw is not None:
            chk('Contingency: Forecast remaining = Contingency − Total drawdown', remain, round(carried-draw,2))
    # §2.2 Development Summary: Approved-Variations column nets to ~0 (variations funded from contingency)
    ds=find(lambda t: t.cell(0,2).text.strip().lower().startswith('approved variation') and any(t.cell(r,0).text.strip().lower().startswith('total') for r in range(len(t.rows))))
    if ds:
        tot=next((r for r in range(len(ds.rows)) if ds.cell(r,0).text.strip().lower().startswith('total')), None)
        if tot is not None:
            col=sum((_money(ds.cell(r,2).text) or 0) for r in range(1,tot))
            chk('Dev Summary: Σ Approved Variations column = displayed total (funded from contingency)', round(col,0), _money(ds.cell(tot,2).text), tol=2.0)
    return checks

def write_qa(docx_path, checks):
    out=docx_path.replace('.docx','_QA.txt')
    lines=['QA CHECKS — '+os.path.basename(docx_path),'='*60,'']
    npass=sum(1 for c in checks if c[1])
    for name,ok,got,expect,delta in checks:
        flag='PASS' if ok else '**FAIL**'
        gv='' if got is None else f'{got:,.2f}'; ev='' if expect is None else f'{expect:,.2f}'
        d='' if delta in (None,) else f'  (Δ {delta:+,.2f})'
        lines.append(f'[{flag}] {name}\n        got {gv}  vs expected {ev}{d}')
    lines.append(''); lines.append(f'{npass}/{len(checks)} checks passed.')
    if npass<len(checks): lines.append('REVIEW REQUIRED — carry-forward or internal-consistency error detected.')
    open(out,'w').write('\n'.join(lines))
    print('QA:', f'{npass}/{len(checks)} passed', '->', os.path.basename(out))
    for name,ok,*_ in checks:
        if not ok: print('   FAIL:', name)


# --- BDM standing drafting rule (CN-pending): risk-rating colour coding ---
# LOW = green 1A4A2A, MEDIUM = amber 9D7B5B, HIGH = red B91C1C (bold).
RISK_RATING_COLOURS = {"LOW": "1A4A2A", "MEDIUM": "9D7B5B", "HIGH": "B91C1C"}
def colour_risk_ratings(d):
    """Colour every risk-rating cell by its value per the BDM standing rule.
    Whole-cell exact match only (LOW/MEDIUM/HIGH, any case), so prose that merely
    contains the words is left untouched. Covers the Risk Level / Rating columns and
    any Likelihood/Impact cells that use the same scale."""
    from docx.shared import RGBColor
    n = 0
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                key = cell.text.strip().upper()
                if key in RISK_RATING_COLOURS:
                    rgb = RGBColor.from_string(RISK_RATING_COLOURS[key])
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.text.strip():
                                r.font.color.rgb = rgb
                                r.font.bold = True
                    n += 1
    return n

def build(args):
    cfg=json.load(open(args.config))
    d=docx.Document(args.template)

    TOK=cfg.get('global_replacements',{})
    def repl_runs(p):
        for r in p.runs:
            for k,v in TOK.items():
                if k in r.text: r.text=r.text.replace(k,v)
    for p in d.paragraphs: repl_runs(p)
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs: repl_runs(p)
    fallback_token_replace(d, TOK)

    apply_inserts(d, cfg)
    # delete rows (descending) before cell fills — for variable-length body tables
    for ti_s, rows in cfg.get('trim_rows',{}).items():
        ti=int(ti_s)
        for r in sorted(rows, reverse=True):
            if r < len(d.tables[ti].rows): del_row(d, ti, r)

    for ti_s, cells in cfg.get('table_cells',{}).items():
        ti=int(ti_s)
        for r,c,text in cells: set_cell(d,ti,r,c,text)
    for idx_s, text in cfg.get('para_set',{}).items():
        setpar(d.paragraphs[int(idx_s)], text)

    tb=cfg.get('trade_breakdown')
    if tb:
        ti=tb['table']; rows=tb['rows']; start=tb.get('start_row',2)
        for i,tr in enumerate(rows):
            r=start+i
            for c,key in enumerate(('desc','sv','prevd','prevpct','cur','curpct','totd','totpct')):
                set_cell(d,ti,r,c,tr[key])
        subr=start+len(rows)
        for c,val in enumerate(tb['subtotal']): set_cell(d,ti,subr,c,val)
        while len(d.tables[ti].rows)>subr+1: del_row(d,ti,len(d.tables[ti].rows)-1)
        bold_shade_row(d,ti,subr)

    va=cfg.get('variations')
    if va:
        ti=va['table']; rows=va['rows']; start=va.get('start_row',2)
        for i,vr in enumerate(rows):
            for c,val in enumerate(vr): set_cell(d,ti,start+i,c,val)
        subr=start+len(rows)
        for c,val in enumerate(va['subtotal']): set_cell(d,ti,subr,c,val)
        while len(d.tables[ti].rows)>subr+1: del_row(d,ti,len(d.tables[ti].rows)-1)
        bold_shade_row(d,ti,subr)

    cf=cfg.get('cashflow')
    if cf and cf.get('rows'):
        ti=cf['table']; start=cf.get('start_row',1)
        for i,row in enumerate(cf['rows']):
            for c,val in enumerate(row): set_cell(d,ti,start+i,c,val)
        while len(d.tables[ti].rows)>start+len(cf['rows']): del_row(d,ti,len(d.tables[ti].rows)-1)

    co=cfg.get('consultants')
    if co:
        ti=co['table']; rows=co['rows']; start=co.get('start_row',1); t=d.tables[ti]
        for i,cr in enumerate(rows):
            if start+i >= len(t.rows): t._tbl.append(copy.deepcopy(t.rows[start]._tr))
            for c,val in enumerate(cr): set_cell(d,ti,start+i,c,val)
        while len(d.tables[ti].rows) > start+len(rows): del_row(d,ti,len(d.tables[ti].rows)-1)

    for note in cfg.get('remove_notes',[]):
        for p in list(d.paragraphs):
            if note in p.text: p._p.getparent().remove(p._p)
    for keepspec in cfg.get('trim_notes',[]):
        for p in d.paragraphs:
            if keepspec['match'] in p.text: setpar(p, keepspec['keep'])

    ed=cfg.get('expired_dates')
    if ed:
        report=parse_date(ed['report_date'])
        for spec in ed.get('targets',[]):
            ti=spec['table']; col=spec['col']; t=d.tables[ti]
            for ri in range(1,len(t.rows)):
                cells=t.rows[ri].cells
                if col>=len(cells): continue
                txt=cells[col].text.strip(); dt=parse_date(txt)
                if txt and (txt.lower() in ('not received','-','—') or (dt and report and dt<report)):
                    set_cell(d, ti, ri, col, txt, red=True)

    apply_row_format(d, cfg)
    apply_bold_shade(d, cfg)
    ph=cfg.get('photos')
    if ph and ph.get('captions'):
        caps=ph['captions']; idx=0
        for ti in ph.get('caption_tables',[24,25]):
            if ti>=len(d.tables): continue
            t=d.tables[ti]
            for r in range(len(t.rows)):
                for c in range(len(t.columns)):
                    for p in t.cell(r,c).paragraphs:
                        if ('Caption' in p.text or '[' in p.text) and idx<len(caps):
                            setpar(p, caps[idx]); idx+=1; break

    clone_2col_table(d, cfg)
    embed_appendices(d, cfg, args.claimdir)
    add_additional_appendices(d, cfg, args.claimdir)
    clean_appendix_markers(d)
    justify_body(d)
    colour_risk_ratings(d)

    draft=cfg.get('draft','DRAFT — FOR SENIOR PM / QS REVIEW — NOT FOR ISSUE')
    p=d.tables[0].cell(0,0).add_paragraph(); run=p.add_run(draft)
    run.bold=True; run.font.color.rgb=RGBColor(0xC0,0,0); run.font.size=Pt(11)

    cleanup_blank_paras(d)

    tmp_docx=args.out.replace('.docx','_t1.docx')
    d.save(tmp_docx)
    final=phase2(tmp_docx, args, cfg)
    try: os.replace(final, args.out)
    except OSError:
        import shutil; shutil.copy(final, args.out)
    for tmp in (tmp_docx, final):
        try: os.remove(tmp)
        except OSError: pass
    print('BUILT', args.out)
    if args.pdf: render_pdf(args.out, cfg, args.claimdir)
    try:
        write_qa(args.out, qa_checks(args.out, cfg))
    except Exception as e:
        print('QA skipped:', e)

def phase2(src, args, cfg):
    out=src.replace('_t1.docx','_t2.docx')
    zin=zipfile.ZipFile(src)
    doc=zin.read('word/document.xml').decode('utf8')
    rels=zin.read('word/_rels/document.xml.rels').decode('utf8')
    extra_media=[]
    ph=cfg.get('photos')
    if ph and args.photos:
        photo_files=extract_photos(args.photos, ph.get('count',16))
        maxrid=max(int(x) for x in re.findall(r'Id="rId(\d+)"',rels))
        newrels=[]; rids=[]
        for k,pf in enumerate(photo_files):
            rid='rId%d'%(maxrid+1+k); rids.append(rid)
            newrels.append('<Relationship Id="%s" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/wphoto_%02d.jpg"/>'%(rid,k))
            extra_media.append(('word/media/wphoto_%02d.jpg'%k, pf))
        rels=rels.replace('</Relationships>',''.join(newrels)+'</Relationships>')
        cnt=[0]
        def rp(m):
            seg=m.group(0)
            if 'name="site-photo"' in seg and cnt[0]<len(rids):
                seg=re.sub(r'(<a:blip[^>]*r:embed=")rId\d+(")', r'\g<1>%s\g<2>'%rids[cnt[0]], seg, count=1); cnt[0]+=1
            return seg
        doc=re.sub(r'<w:drawing\b.*?</w:drawing>', rp, doc, flags=re.S)
    # remove leftover appendix-B sample images
    doc=re.sub(r'<w:drawing\b.*?</w:drawing>', lambda m: '' if (re.search(r'name="appx-b-\d+"',m.group(0)) or 'name="appx-23"' in m.group(0)) else m.group(0), doc, flags=re.S)
    # XML-level token pass — catches tokens in nested tables / split across runs that python-docx missed
    for k,v in cfg.get('global_replacements',{}).items():
        doc=doc.replace(k,v)
    chart=None
    cf=cfg.get('cashflow')
    if cf and cf.get('chart_cats'):
        chart=zin.read('word/charts/chart1.xml').decode('utf8')
        def strc(v): return '<c:strCache><c:ptCount val="%d"/>%s</c:strCache>'%(len(v),''.join('<c:pt idx="%d"><c:v>%s</c:v></c:pt>'%(i,x) for i,x in enumerate(v)))
        def numc(v): return '<c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="%d"/>%s</c:numCache>'%(len(v),''.join('<c:pt idx="%d"><c:v>%s</c:v></c:pt>'%(i,x) for i,x in enumerate(v)))
        chart=re.sub(r'<c:cat>.*?</c:cat>', lambda m: re.sub(r'<c:strCache>.*?</c:strCache>',strc(cf['chart_cats']),m.group(0),flags=re.S), chart, flags=re.S)
        it=iter([numc(cf['chart_forecast']),numc(cf['chart_actual'])])
        chart=re.sub(r'<c:numCache>.*?</c:numCache>', lambda m: next(it), chart, count=2, flags=re.S)
    htok=cfg.get('header_replacements', cfg.get('global_replacements',{}))
    zout=zipfile.ZipFile(out,'w',zipfile.ZIP_DEFLATED)
    for it in zin.infolist():
        n=it.filename; data=zin.read(n)
        if n=='word/document.xml': data=doc.encode('utf8')
        elif n=='word/_rels/document.xml.rels': data=rels.encode('utf8')
        elif n=='word/charts/chart1.xml' and chart is not None: data=chart.encode('utf8')
        elif re.match(r'word/(header|footer)\d*\.xml$',n):
            x=data.decode('utf8')
            for k,v in htok.items(): x=x.replace(k,v)
            data=x.encode('utf8')
        zout.writestr(it,data)
    for tgt,pf in extra_media: zout.writestr(tgt, open(pf,'rb').read())
    zout.close(); zin.close()
    return out

def extract_photos(pdf_path, count):
    if os.path.isdir(pdf_path):
        exts=('.jpg','.jpeg','.png')
        return sorted(f for f in glob.glob(os.path.join(pdf_path,'*')) if f.lower().endswith(exts))[:count]
    import fitz
    from PIL import Image, ImageChops
    import numpy as np
    doc=fitz.open(pdf_path); outdir=tempfile.mkdtemp()
    def autocrop(im):
        bbox=ImageChops.difference(im,Image.new('RGB',im.size,(255,255,255))).getbbox()
        if bbox:
            l,t,r,b=bbox; return im.crop((max(l-3,0),max(t-3,0),min(r+3,im.size[0]),min(b+3,im.size[1])))
        return im
    def split(body):
        a=np.asarray(body.convert('L')); wr=a.min(axis=1)>235; H=len(wr); bands=[]; i=0
        while i<H:
            if wr[i]:
                j=i
                while j<H and wr[j]: j+=1
                bands.append((i,j,j-i)); i=j
            else: i+=1
        mids=[b for b in bands if b[0]>H*0.18 and b[1]<H*0.82 and b[2]>8]
        if mids:
            sep=max(mids,key=lambda b:b[2]); cut=(sep[0]+sep[1])//2
            return [autocrop(body.crop((0,0,body.size[0],cut))), autocrop(body.crop((0,cut,body.size[0],body.size[1])))]
        return [autocrop(body)]
    n=0; files=[]
    for pno in range(1, doc.page_count):
        pix=doc[pno].get_pixmap(matrix=fitz.Matrix(2.4,2.4))
        im=Image.frombytes('RGB',[pix.width,pix.height],pix.samples); W,H=im.size
        for part in split(im.crop((0,0,W,int(H*0.94)))):
            w,h=part.size
            if w>250 and h>200 and 0.4<h/w<2.2 and n<count:
                fp=os.path.join(outdir,'ph_%02d.jpg'%n); part.save(fp,quality=88); files.append(fp); n+=1
        if n>=count: break
    return files

def render_pdf(docx_path, cfg=None, claimdir=None):
    work=tempfile.mkdtemp()
    subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',work,docx_path],
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=300)
    src=os.path.join(work, os.path.splitext(os.path.basename(docx_path))[0]+'.pdf')
    dst=docx_path.replace('.docx','.pdf')
    if not os.path.exists(src): return
    ca=(cfg or {}).get('compile_appendices')
    if not ca:
        import shutil; shutil.copy(src,dst); print('RENDERED', dst); return
    from pypdf import PdfReader, PdfWriter
    body=PdfReader(src)
    def find(lbl):
        for i in range(len(body.pages)):
            if lbl in (body.pages[i].extract_text() or '') and i>5: return i
        return None
    B=find('Appendix B:'); C=find('Appendix C:'); w=PdfWriter()
    upto=(B+1) if B is not None else len(body.pages)
    for i in range(0,upto): w.add_page(body.pages[i])
    for f in ca.get('appendix_b',[]):
        p=f if os.path.isabs(f) else os.path.join(claimdir or '',f)
        if os.path.exists(p):
            for pg in PdfReader(p).pages: w.add_page(pg)
    if C is not None:
        for i in range(B+1,C+1): w.add_page(body.pages[i])
        note=ca.get('appendix_c_note')
        if note:
            for pg in PdfReader(_statdec_pdf(work, note)).pages: w.add_page(pg)
        for i in range(C+1,len(body.pages)): w.add_page(body.pages[i])
    out=docx_path.replace('.docx',' (with Appendices).pdf')
    with open(out,'wb') as fh: w.write(fh)
    import shutil; shutil.copy(src,dst)
    print('RENDERED', dst); print('COMPILED', out, len(w.pages),'pages')

def _statdec_pdf(work, note):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    import textwrap
    fp=os.path.join(work,'statdec.pdf'); c=canvas.Canvas(fp,pagesize=A4); W,H=A4
    c.setFont('Helvetica-Bold',13); c.drawString(25*mm,H-45*mm,'Statutory Declaration (Building)')
    c.setFont('Helvetica',11); y=H-60*mm
    for ln in textwrap.wrap(note,90): c.drawString(25*mm,y,ln); y-=7*mm
    c.showPage(); c.save(); return fp

if __name__=='__main__':
    ap=argparse.ArgumentParser()
    ap.add_argument('--template',required=True); ap.add_argument('--config',required=True)
    ap.add_argument('--photos'); ap.add_argument('--claimdir'); ap.add_argument('--out',required=True)
    ap.add_argument('--pdf',action='store_true')
    build(ap.parse_args())
