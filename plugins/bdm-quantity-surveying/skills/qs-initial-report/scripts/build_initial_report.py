#!/usr/bin/env python3
"""
build_initial_report.py - BDM Form 424 QS Initial Report draft builder.

Division of labour (BDM house pattern): Claude reads the financier brief and the project's
"Information Received" folder, makes the QS judgements, and writes a JSON content map. THIS
script does the deterministic build from that map:
  - global token replacement (body + text boxes/content controls + headers/footers)
  - named table-cell fill
  - narrative paragraph replacement by anchor
  - TWO-CATEGORY flag system: retire funder tags; relabel the template's generic prompts to
    BDM (bronze internal input/consideration prompts); add FINANCIER residual-risk flags in a
    distinct dark-red callout
  - appendices restructured to the Form 425 convention (Heading 2 "Appendix X:") with the
    received PDF extracts embedded one page per Word page (rasterised); notes for not-received
  - DRAFT marker; save .docx
It never invents content - only what the content map provides.

Usage:
  python build_initial_report.py --template <424.docx> --config <map.json> \
      --infodir "<...>/00_Information Received" --out "<...>/DRAFT_...v0.1.docx"
Requires: python-docx, pdftoppm (poppler-utils) for appendix embedding.
"""
import argparse, json, copy, sys, re, os, glob, tempfile, subprocess
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Emu
from docx.enum.text import WD_BREAK

FUNDER = {"ANZ", "NAB", "CRER", "WESTPAC"}
KEEP = {"COMPLETE", "BDM", "FINANCIER", "REVIEW"}
FIN_FILL, FIN_BAR, FIN_LABEL = "F8ECEC", "A33A3A", "A33A3A"   # dark-red Financier callout

# ---------- text helpers ----------
def para_text(p):
    return "".join(n.text for n in p._p.iter(qn('w:t')) if n.text)

def is_flag(p):
    return para_text(p).strip().upper().startswith("FLAG")

def flag_tags(p):
    m = re.match(r"FLAG\s*[—-]\s*([^.]*)", para_text(p).strip())
    if not m:
        return set()
    return {t.strip().upper() for t in re.split(r"[\s/\-—,]+", m.group(1)) if t.strip()}

def set_para_text(p, new_text):
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""

def iter_body_paragraphs(doc):
    yield from doc.paragraphs
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                yield from cell.paragraphs

# ---------- global token replacement (XML level) ----------
def _para_replace_xml(p_el, keys, repl):
    ts = p_el.findall('.//' + qn('w:t'))
    if not ts:
        return
    full = "".join(t.text or "" for t in ts)
    if "[" not in full or full.strip().upper().startswith("FLAG"):
        return
    new = full
    for k in keys:
        if k in new:
            new = new.replace(k, repl[k])
    if new == full:
        return
    ts[0].text = new
    for t in ts[1:]:
        t.text = ""

def apply_global(doc, repl):
    keys = sorted(repl.keys(), key=len, reverse=True)
    roots = [doc.element.body]
    for sec in doc.sections:
        for hf in (sec.first_page_header, sec.header, sec.even_page_header,
                   sec.first_page_footer, sec.footer, sec.even_page_footer):
            try:
                roots.append(hf._element)
            except Exception:
                pass
    for root in roots:
        for p_el in root.iter(qn('w:p')):
            _para_replace_xml(p_el, keys, repl)

# ---------- table cells ----------
def set_table_cells(doc, table_cells):
    for spec in table_cells:
        cell = doc.tables[spec["table"]].cell(spec["row"], spec["col"])
        ps = cell.paragraphs
        if ps:
            set_para_text(ps[0], spec["text"])
            for extra in ps[1:]:
                extra._p.getparent().remove(extra._p)
        else:
            cell.text = spec["text"]

# ---------- narrative paragraphs ----------
def replace_paragraphs(doc, para_set):
    for spec in para_set:
        anchor = spec["anchor"]
        for p in iter_body_paragraphs(doc):
            if is_flag(p):
                continue
            if anchor in "".join(r.text for r in p.runs):
                set_para_text(p, spec["text"])
                break

# ---------- flags (two-category) ----------
def recategorise_flags(doc, retire_funder=True, relabel_complete_to_bdm=True):
    """Retire funder-tagged flags (ANZ/NAB/CRER/WESTPAC); relabel generic COMPLETE prompts to
    BDM (internal input/consideration prompts; bronze style unchanged)."""
    removed = []
    for p in list(doc.paragraphs):
        if not is_flag(p):
            continue
        tags = flag_tags(p)
        if not tags:
            continue
        if retire_funder and (tags & FUNDER) and not (tags & KEEP):
            removed.append(para_text(p)[:70])
            p._p.getparent().remove(p._p)
            continue
        if relabel_complete_to_bdm and "COMPLETE" in tags and p.runs:
            p.runs[0].text = p.runs[0].text.replace("COMPLETE", "BDM")
    return removed

def _recolour_flag(p_el, fill, bar, label_color):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        return
    shd = pPr.find(qn('w:shd'))
    if shd is not None:
        shd.set(qn('w:fill'), fill)
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is not None:
        left = pBdr.find(qn('w:left'))
        if left is not None:
            left.set(qn('w:color'), bar)
    r0 = p_el.find(qn('w:r'))
    if r0 is not None:
        rpr = r0.find(qn('w:rPr'))
        if rpr is not None:
            col = rpr.find(qn('w:color'))
            if col is not None:
                col.set(qn('w:val'), label_color)

def add_flags(doc, additions):
    """Insert a flag after the first heading matching anchor_heading, cloning the house bronze
    callout. kind='BDM' (bronze) or 'FINANCIER' (recoloured dark red)."""
    template_flag = None
    for p in doc.paragraphs:
        if is_flag(p):
            template_flag = p
            break
    if template_flag is None:
        return
    for spec in additions:
        anchor = spec["anchor_heading"]
        kind = spec.get("kind", "BDM").upper()
        label = spec.get("label", kind)
        body = spec["body"]
        target = None
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith("Heading") and anchor in para_text(p):
                target = p
                break
        if target is None:
            continue
        new_p = copy.deepcopy(template_flag._p)
        tmp = template_flag.__class__(new_p, template_flag._parent)
        runs = tmp.runs
        if len(runs) >= 2:
            runs[0].text = "FLAG — " + label + ".  "
            runs[1].text = body
            for r in runs[2:]:
                r.text = ""
        else:
            set_para_text(tmp, "FLAG — " + label + ".  " + body)
        if kind == "FINANCIER":
            _recolour_flag(new_p, FIN_FILL, FIN_BAR, FIN_LABEL)
        target._p.addnext(new_p)

# ---------- appendices (425 convention + PDF embed) ----------
def _expand_pages(pages, total):
    if pages in (None, "", "all"):
        return list(range(1, total + 1))
    out = []
    for part in str(pages).split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-")
            out += list(range(int(a), int(b) + 1))
        elif part:
            out.append(int(part))
    return [p for p in out if 1 <= p <= total]

def _rasterise(pdf, pages, tmpdir, dpi=150):
    """Rasterise the requested 1-based pages of pdf to PNGs; return ordered list of paths."""
    try:
        info = subprocess.run(["pdfinfo", pdf], capture_output=True, text=True)
        total = int(re.search(r"Pages:\s+(\d+)", info.stdout).group(1))
    except Exception:
        total = 9999
    wanted = _expand_pages(pages, total)
    paths = []
    for i, pg in enumerate(wanted):
        prefix = os.path.join(tmpdir, "ap_%d_%d" % (abs(hash(pdf)) % 100000, i))
        subprocess.run(["pdftoppm", "-png", "-r", str(dpi), "-f", str(pg), "-l", str(pg),
                        pdf, prefix], capture_output=True)
        got = sorted(glob.glob(prefix + "*.png"))
        paths += got
    return paths

def insert_appendices(doc, specs, infodir, tmpdir):
    """Replace the plain 'Appendix X - Title' list with Form 425-style Heading 2 sub-headings,
    embedding received PDF extracts one page per Word page, or a note if not received."""
    if not specs:
        return {"embedded_pages": 0, "appendices": 0}
    # remove the existing plain appendix list lines
    for p in list(doc.paragraphs):
        if p.style and p.style.name == "Normal" and re.match(r"Appendix [A-I]\s*[—-]", p.text.strip()):
            p._p.getparent().remove(p._p)
    sec = doc.sections[0]
    width = Emu(sec.page_width - sec.left_margin - sec.right_margin)
    embedded = 0
    for s in specs:
        doc.add_paragraph("Appendix %s: %s" % (s["letter"], s["title"]), style="Heading 2")
        if s.get("note") and not s.get("sources"):
            doc.add_paragraph(s["note"])
            continue
        any_page = False
        for src in s.get("sources", []):
            pdf = os.path.join(infodir, src["pdf"])
            if not os.path.isfile(pdf):
                doc.add_paragraph("[Source not found: %s]" % src["pdf"])
                continue
            for img in _rasterise(pdf, src.get("pages", "all"), tmpdir):
                par = doc.add_paragraph()
                run = par.add_run()
                run.add_break(WD_BREAK.PAGE)
                run.add_picture(img, width=width)
                embedded += 1
                any_page = True
        if not any_page and s.get("note"):
            doc.add_paragraph(s["note"])
    return {"embedded_pages": embedded, "appendices": len(specs)}

# ---------- draft marker ----------
def mark_draft(doc, cover_table, draft_line):
    try:
        cell = doc.tables[cover_table].cell(0, 0)
        marker = cell.paragraphs[0].insert_paragraph_before("")
        run = marker.add_run(draft_line)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    except Exception as e:
        print("  (draft marker note: %s)" % e, file=sys.stderr)

# ---------- main ----------

# --- BDM standing drafting rule (CN-pending): risk-rating colour coding ---
# LOW = green 1A4A2A, MEDIUM = amber 9D7B5B, HIGH = red B91C1C (bold).
RISK_RATING_COLOURS = {"LOW": "1A4A2A", "MEDIUM": "9D7B5B", "HIGH": "B91C1C"}
def colour_risk_ratings(d):
    """Colour every risk-rating cell by its value per the BDM standing rule.
    Whole-cell exact match only (LOW/MEDIUM/HIGH, any case), so prose that merely
    contains the words is left untouched. Covers the Risk Level / Rating columns and
    any Likelihood/Impact cells that use the same scale."""
    from docx.shared import RGBColor
    n = 0
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                key = cell.text.strip().upper()
                if key in RISK_RATING_COLOURS:
                    rgb = RGBColor.from_string(RISK_RATING_COLOURS[key])
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.text.strip():
                                r.font.color.rgb = rgb
                                r.font.bold = True
                    n += 1
    return n

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--config", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--infodir", default="", help="Information Received folder (appendix PDFs)")
    a = ap.parse_args()

    cfg = json.load(open(a.config, encoding="utf-8"))
    doc = Document(a.template)
    tmpdir = tempfile.mkdtemp(prefix="qsir_")

    apply_global(doc, cfg.get("global_replacements", {}))
    set_table_cells(doc, cfg.get("table_cells", []))
    replace_paragraphs(doc, cfg.get("para_set", []))
    fl = cfg.get("flags", {})
    removed = recategorise_flags(doc, fl.get("retire_funder_tags", True),
                                 fl.get("relabel_complete_to_bdm", True))
    add_flags(doc, fl.get("add", []))
    apx = insert_appendices(doc, cfg.get("appendices", []), a.infodir or cfg.get("infodir", ""), tmpdir)

    dm = cfg.get("draft", {})
    default_draft = "DRAFT - FOR SENIOR PM / QS REVIEW - NOT FOR ISSUE"
    if dm.get("mark", True):
        mark_draft(doc, dm.get("cover_table", 0), dm.get("line", default_draft))

    colour_risk_ratings(doc)
    doc.save(a.out)
    print(json.dumps({
        "saved": a.out,
        "funder_flags_retired": len(removed),
        "appendices": apx["appendices"],
        "appendix_pages_embedded": apx["embedded_pages"],
        "tables": len(doc.tables),
    }, indent=2))

if __name__ == "__main__":
    main()
