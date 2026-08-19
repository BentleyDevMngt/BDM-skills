#!/usr/bin/env python3
"""
Build a BDM Site Inspection Record (Form 331) from a JSON config + stitched photos.

USAGE
-----
    python3 build_report.py config.json output.docx

CONFIG SHAPE (see example_config.json)
--------------------------------------
{
  "template": "/abs/path/to/331-Site_Inspection_Record_Rx_YYYY-MM.docx",  # optional; auto-found if omitted
  "particulars": {"ProjectName": "...", "ProjectNumber": "...",
                   "InspectionDate": "YYYY-MM-DD", "InspectionTime": "...",
                   "Weather": "...", "StageOfWorks": "..."},
  "attendees": [["Name", "Company / Role"], ...],
  "observations": [["1.0","BACKGROUND","",true], ["1.1","text","Owner",false], ...],
  "photos_dir": "/abs/dir/with/photo1.jpg...",
  "captions": ["Photo 1 - ...", "Photo 2 - ...", ...],
  "options": {"photos_per_page": 4, "signature": false}
}
Observation rows: [item, description, action, is_section_header].
"""
import sys, os, json, glob, copy
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0F, 0x17, 0x21)
# photos_per_page -> (columns, image width cm)
LAYOUT = {1: (1, 12.0), 2: (2, 7.2), 4: (2, 5.2)}

TEMPLATE_GLOB = ("BDM TEMPLATES/Working Copy/300 Project Management - Contract Delivery/"
                 "331-Site_Inspection_Record_R*_*.docx")

def find_template():
    # search upward from common mount roots
    for root in ("/sessions", os.path.expanduser("~"), "/"):
        hits = glob.glob(os.path.join(root, "**", TEMPLATE_GLOB), recursive=True)
        hits = [h for h in hits if "_Superseded" not in h]
        if hits:
            return sorted(hits)[-1]   # latest revision by name
    raise SystemExit("Form 331 template not found - pass 'template' in config.")

def set_tc_text(tc, text, bold=False):
    p = tc.findall(qn('w:p'))[0]
    r = p.findall(qn('w:r'))[0]
    for t in r.findall(qn('w:t')):
        r.remove(t)
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    if bold and b is None: rPr.append(OxmlElement('w:b'))
    if not bold and b is not None: rPr.remove(b)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)

def main(cfg_path, out_path):
    cfg = json.load(open(cfg_path, encoding='utf-8'))
    template = cfg.get("template") or find_template()
    doc = Document(template)
    opts = cfg.get("options", {})
    ppp = opts.get("photos_per_page", 4)
    cols, wcm = LAYOUT.get(ppp, LAYOUT[4])

    # ---- Table 0: particulars (replace {{Field}} placeholders) ----
    part = cfg.get("particulars", {})
    for row in doc.tables[0].rows:
        vc = row.cells[1]; txt = vc.text.strip()
        m = txt.strip("{}").split(":")[0] if txt.startswith("{{") else None
        if m and m in part:
            for para in vc.paragraphs:
                for run in para.runs:
                    if run.text.strip() == txt:
                        run.text = part[m]

    # ---- Table 1: attendees (fill, then delete blank rows) ----
    att = cfg.get("attendees", [])
    t1 = doc.tables[1]
    def setcell(cell, text):
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for e in p.runs[1:]: e.text = ''
        else:
            run = p.add_run(text); run.font.name = 'Aptos'; run.font.size = Pt(10)
            run.font.color.rgb = NAVY
    body = t1.rows[1:]
    for i, brow in enumerate(body):
        if i < len(att):
            setcell(brow.cells[0], att[i][0]); setcell(brow.cells[1], att[i][1])
        else:
            t1._tbl.remove(brow._tr)

    # ---- Table 2: observations (clone item-row template) ----
    obs = cfg.get("observations", [])
    t2 = doc.tables[2]; tbl = t2._tbl
    item_template = copy.deepcopy(t2.rows[2]._tr)
    for r in list(t2.rows)[1:]:
        tbl.remove(r._tr)
    for row in obs:
        item, desc, action, is_sec = (row + ["", "", "", False])[:4]
        ntr = copy.deepcopy(item_template); tcs = ntr.findall(qn('w:tc'))
        set_tc_text(tcs[0], str(item), bold=is_sec)
        set_tc_text(tcs[1], str(desc), bold=is_sec)
        set_tc_text(tcs[2], str(action), bold=False)
        tbl.append(ntr)

    # ---- locate PHOTOGRAPHS heading + placeholder ----
    photo_heading = placeholder = None
    for p in doc.paragraphs:
        if p.text.strip() == 'PHOTOGRAPHS': photo_heading = p
        if 'Insert site photographs' in p.text: placeholder = p

    # ---- optional signature removal ----
    if not opts.get("signature", False):
        rm = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith('BDM Site Representative'):
                rm.append(p)
            else:
                bd = p._p.find('.//' + qn('w:pBdr'))
                if bd is not None and t and set(t) <= set(' '):
                    rm.append(p)   # the spaces-only signature rule line
        for p in rm:
            p._p.getparent().remove(p._p)

    # ---- placeholder out, page break before heading ----
    if placeholder is not None:
        placeholder._p.getparent().remove(placeholder._p)
    pPr = photo_heading._p.get_or_add_pPr()
    pPr.insert(0, OxmlElement('w:pageBreakBefore'))

    # ---- photo grid right after the heading ----
    photos_dir = cfg.get("photos_dir")
    captions = cfg.get("captions", [])
    photo_files = sorted(glob.glob(os.path.join(photos_dir, "photo*.jpg")),
                         key=lambda f: int(''.join(filter(str.isdigit, os.path.basename(f))) or 0)) if photos_dir else []
    if photo_files:
        ptable = doc.add_table(rows=0, cols=cols); ptable.autofit = False
        for i, pf in enumerate(photo_files):
            if i % cols == 0:
                row = ptable.add_row()
            cell = row.cells[i % cols]
            ip = cell.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.add_run().add_picture(pf, width=Cm(wcm))
            cap = cell.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = captions[i] if i < len(captions) else ""
            cr = cap.add_run(text); cr.font.name = 'Aptos'; cr.font.size = Pt(8)
            cr.font.italic = True; cr.font.color.rgb = NAVY
        photo_heading._p.addnext(ptable._tbl)

    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Template: {template}")
    print(f"Photos: {len(photo_files)} at {ppp}/page (cols={cols}, w={wcm}cm) | signature={opts.get('signature', False)}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__); sys.exit(1)
    main(sys.argv[1], sys.argv[2])
