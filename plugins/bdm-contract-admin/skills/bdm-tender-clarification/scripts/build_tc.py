#!/usr/bin/env python3
"""
build_tc.py — build a BDM Tender Clarification (Form 218) Word doc from a config JSON.

Company standard (bdm-tender-clarification): clone the project's latest TC (a full
Form 218 carrying the brand + particulars), then fill it in. Handles any combination
of: tender-close extension, consolidated RFI responses, supplementary attachments.
Renumbers the Submission / Acknowledgement sections automatically. Optionally swaps
the author block + signature.

Usage:
    python3 build_tc.py config.json out.docx

See example_config.json for the config shape. Reference details (drawing nos,
revisions, dates) must already be in the config — read them off the source
documents; never invent them.
"""
import sys, json, copy
from docx import Document
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

W_T = qn('w:t'); W_P = qn('w:p'); W_TBL = qn('w:tbl')

import re
HEAD_RE = re.compile(r'^\s*\d+\.0\s+(.*\S)\s*$')


def para_text(p_el):
    return ''.join(t.text or '' for t in p_el.iter(W_T))


def set_para_text(p, text):
    """Set a python-docx Paragraph's text, preserving the first run's formatting."""
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def heading_iter(body):
    for el in body:
        if el.tag == W_P:
            m = HEAD_RE.match(para_text(el))
            if m:
                yield el, m.group(1).upper()


def find_heading(body, keyword):
    for el, title in heading_iter(body):
        if keyword.upper() in title:
            return el
    return None


def next_p(el):
    nx = el.getnext()
    while nx is not None and nx.tag != W_P:
        nx = nx.getnext()
    return nx


def next_tbl(el):
    nx = el.getnext()
    while nx is not None and nx.tag != W_TBL:
        nx = nx.getnext()
    return nx


def remove(el):
    if el is not None and el.getparent() is not None:
        el.getparent().remove(el)


def update_particulars(doc, P):
    label_map = {
        "PROJECT": P.get("project"),
        "PROJECT NUMBER": P.get("project_number"),
        "PROJECT MANAGER": P.get("project_manager"),
        "ISSUED TO": P.get("issued_to"),
        "CLARIFICATION NUMBER": P.get("clarification_number"),
        "ISSUE DATE": P.get("issue_date"),
    }
    for row in doc.tables[0].rows:
        label = row.cells[0].text.strip().upper()
        if label in label_map and label_map[label] is not None:
            set_para_text(row.cells[1].paragraphs[0], label_map[label])


def set_subtitle(doc, text):
    if text is None:
        return
    body = doc.element.body
    h1 = find_heading(body, "PROJECT")
    if h1 is None:
        return
    prev = h1.getprevious()
    while prev is not None:
        if prev.tag == W_P and para_text(prev).strip():
            set_para_text(Paragraph(prev, None), text)
            return
        prev = prev.getprevious()


def set_section_body(doc, keyword, text):
    if text is None:
        return
    h = find_heading(doc.element.body, keyword)
    if h is None:
        return
    p = next_p(h)
    if p is not None:
        set_para_text(Paragraph(p, None), text)


def fill_rfi_table(doc, items):
    h = find_heading(doc.element.body, "RFI")
    if h is None:
        return False
    tbl_el = next_tbl(h)
    if tbl_el is None:
        return False
    table = next(t for t in doc.tables if t._tbl is tbl_el)
    data_rows = table.rows[1:]
    template_tr = data_rows[-1]._tr
    have, need = len(data_rows), len(items)
    for _ in range(need - have):
        template_tr.addnext(copy.deepcopy(template_tr))
    if need < have:
        for r in data_rows[need:]:
            remove(r._tr)
    table = next(t for t in doc.tables if t._tbl is tbl_el)
    for row, it in zip(table.rows[1:], items):
        vals = [it.get("item", ""), it.get("subject", ""), it.get("query", ""),
                it.get("response", ""), it.get("source", "")]
        for cell, v in zip(row.cells, vals):
            set_para_text(cell.paragraphs[0], v)
    return True


def remove_section(doc, keyword):
    h = find_heading(doc.element.body, keyword)
    if h is None:
        return
    to_del = [h]
    nx = h.getnext()
    while nx is not None:
        if nx.tag == W_P and HEAD_RE.match(para_text(nx)):
            break
        to_del.append(nx)
        nx = nx.getnext()
    for el in to_del:
        remove(el)


def fill_attachments(doc, lines):
    h = find_heading(doc.element.body, "ATTACHMENT")
    if h is None:
        return
    bullets, nx = [], h.getnext()
    while nx is not None and not (nx.tag == W_P and HEAD_RE.match(para_text(nx))):
        if nx.tag == W_P:
            bullets.append(nx)
        nx = nx.getnext()
    if not bullets:
        return
    def bt(s):
        return s if s.lstrip().startswith("•") else "•  " + s
    first = bullets[0]
    set_para_text(Paragraph(first, None), bt(lines[0]))
    for b in bullets[1:]:
        remove(b)
    anchor = first
    for line in lines[1:]:
        new_p = copy.deepcopy(first)
        anchor.addnext(new_p)
        set_para_text(Paragraph(new_p, None), bt(line))
        anchor = new_p


def renumber_tail(doc):
    body = doc.element.body
    after, seen3 = [], False
    for el, title in heading_iter(body):
        if "TENDER CLOSING" in title:
            seen3 = True
            continue
        if seen3:
            after.append(el)
    n = 4
    for el in after:
        m = HEAD_RE.match(para_text(el))
        if m:
            set_para_text(Paragraph(el, None), "%d.0  %s" % (n, m.group(1)))
            n += 1


def set_author_block(doc, author):
    if not author:
        return
    paras = doc.paragraphs
    idx = next((i for i, p in enumerate(paras)
                if p.text.strip().lower().startswith("regards")), None)
    if idx is None:
        return
    rest = [p for p in paras[idx + 1:] if p.text.strip()]
    for p, val in zip(rest, [author.get("name"), author.get("title"), author.get("email")]):
        if val:
            set_para_text(p, val)


def swap_signature(doc, png_path, width_cm):
    paras = doc.paragraphs
    idx = next((i for i, p in enumerate(paras)
                if p.text.strip().lower().startswith("regards")), None)
    if idx is None:
        return False
    for p in paras[idx:idx + 5]:
        if any(True for _ in p._p.iter(qn('w:drawing'))):
            for r in p.runs:
                for dr in r._r.findall(qn('w:drawing')):
                    r._r.remove(dr)
                r.text = ''
            p.add_run().add_picture(png_path, width=Cm(width_cm))
            return True
    return False


def main(argv):
    if len(argv) != 3:
        print(__doc__); return 2
    cfg = json.load(open(argv[1], encoding="utf-8"))
    out = argv[2]
    doc = Document(cfg["source_tc"])

    update_particulars(doc, cfg.get("particulars", {}))
    set_subtitle(doc, cfg.get("subtitle"))
    set_section_body(doc, "CLARIFICATION", cfg.get("clarification"))
    set_section_body(doc, "TENDER CLOSING", cfg.get("closing_date"))

    rfi = cfg.get("rfi_items") or []
    if rfi:
        intro = cfg.get("rfi_intro",
                        "The following table responds to queries raised by Tenderers "
                        "during the tender period. Any drawing referenced in a response "
                        "is attached and forms part of the tender documentation.")
        set_section_body(doc, "RFI", intro)
        if not fill_rfi_table(doc, rfi):
            print("WARN: source TC has no RFI section — clone a full prior TC.", file=sys.stderr)
    else:
        remove_section(doc, "RFI")

    atts = cfg.get("attachments") or []
    if atts:
        fill_attachments(doc, atts)
    else:
        remove_section(doc, "ATTACHMENT")

    renumber_tail(doc)
    set_author_block(doc, cfg.get("author"))
    if cfg.get("signature_png"):
        if not swap_signature(doc, cfg["signature_png"], cfg.get("signature_width_cm", 2.25)):
            print("WARN: could not swap signature image.", file=sys.stderr)

    doc.save(out)
    print("Saved %s" % out)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
